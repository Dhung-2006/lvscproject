import pandas as pd 
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.shared import Cm
from docx2pdf import convert
from pypdf import PdfWriter
import glob
import os
def Cov():    
    try:
        #偵測資料夾xlsx檔案
        excel_files = glob.glob("./excel_to_word/processingData/*.xlsx")
        excel_file = excel_files[0]
        #--------------------------------------------------------------------------------
        #偵測資料夾word檔案
        word_files = glob.glob("./excel_to_word/5.報名表正面.docx")
        word_file = word_files[0]
        #--------------------------------------------------------------------------------
        #為於word & excel 之標題欄位不同處建立dict
        worddic={'身分證號碼':'身分證號',
                '中文姓名':'姓名',
                '出生日期':'出生日期',
                '英文姓名':'英文姓名',
                '原住民傳統姓名並列之羅馬拼音':'身分證上原住民姓名之羅馬拼音',
                '通訊地址':'通訊地址',
                '戶籍地址':'戶籍地址',
                '就讀學校':'報檢人參檢學校',
                '就讀科系':'科系',
                '年級':'年級',
                '班別':'班別',
                '上課別':'部別'
                }
        #讀取主要資料
        df = pd.read_excel(excel_file, sheet_name='Data-全測',skiprows=2)
        rows = df.shape[0]
        #--------------------------------------------------------------------------------
        #讀取&儲存科系資料
        dfs = pd.read_excel(excel_file,sheet_name='代號',usecols=[0])
        text_values = dfs.values
        # 將科系資料加入list
        is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
        only_string = text_values[is_string]
        subject = np.insert(only_string,0,0)
        #---------------------------------------------------------------------------------
        #讀取&儲存班級資料
        dfc  = pd.read_excel(excel_file,sheet_name='代號',usecols=[3])
        text_values = dfc.values
        # 將班級資料加入list
        is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
        only_string = text_values[is_string]
        Class = np.insert(only_string,0,0)
        #---------------------------------------------------------------------------------
        #報檢學校編號與代碼
        dfsc = pd.read_excel(excel_file,sheet_name='代號',usecols=[7,8,9], index_col=0,nrows = 17)
        #---------------------------------------------------------------------------------
        #學制
        dfstu  = pd.read_excel(excel_file,sheet_name='代號',usecols=[18],nrows = 4)
        text_values = dfstu.values
        # 將學制加入lish
        is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
        only_string = text_values[is_string]
        stu = np.insert(only_string,0,0)
        #---------------------------------------------------------------------------------
        #特定對象
        dfsg = pd.read_excel(excel_file,sheet_name='代號',usecols=[20,21],index_col=0, nrows = 8)
        #---------------------------------------------------------------------------------
        #測驗類別
        dftp = pd.read_excel(excel_file,sheet_name='代號',usecols=[14,15],index_col=0, nrows = 13 )
        text_values = dftp.values
        #將測驗類別加入list
        is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
        only_string = text_values[is_string]
        test_type_lst = np.insert(only_string,0,0)
        #---------------------------------------------------------------------------------
        #套印用資料-全測
        df_print = pd.read_excel(excel_file, sheet_name='套印用資料-全測')
        #---------------------------------------------------------------------------------
        #讀取學制
        df_study_type = pd.read_excel(excel_file,sheet_name='代號',usecols=[11,12],index_col=0, nrows = 11)
        text_values = df_study_type.values
        # 將學制加入list
        is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
        only_string = text_values[is_string]
        study_type_list = np.insert(only_string,0,0)
        #---------------------------------------------------------------------------------
        #word讀取 & 填寫
        merger = PdfWriter()  #建立pdf讀寫器
        #讀取每一行的excel主資料
        for i in range(0,rows):
            school_id  ='0'+str(df_print.loc[i,'學號'])
            doc = Document(word_file)
            table =doc.tables[0]     #選定word中第一個表格
            nowcommend = ''   #判斷正在閱讀的儲存格
            testset = set()   #判斷有沒有重複閱讀或輸入的儲存格
            #讀取 行&列
            for idxr,row in enumerate(table.rows) :
                for idxc, cell in enumerate(row.cells) :
                    if cell.text != '':   #該儲存格是否有資料
                        if cell.text not in testset :   #判斷該儲存格是否有重複輸入
                            #填入英文姓名欄位
                            if nowcommend=='英文姓名':
                                if cell.text not in testset:
                                    care = cell.text
                                    paragraph = cell.paragraphs[0]      #為該儲存格 新增段落
                                    paragraph_format = paragraph.paragraph_format #為該儲存格新增文件格式
                                    paragraph_format.line_spacing = Pt(12)   #設定儲存格的行距
                                    for run in paragraph.runs:      
                                        run.clear()
                                    run2 = paragraph.add_run(df.loc[i, worddic[nowcommend]])  #在段落中新增內容
                                    run2.font.size = Pt(12)          # 設置字體大小 
                                    run2.font.name = 'Times New Roman'      #設置英文字型
                                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                                    run1 = paragraph.add_run(care)     #新增第二段文字
                                    run1.font.size = Pt(6)
                                    run1.font.name = '標楷體'#設置中文字型
                                    run1.font.bold = True
                                    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                    testset.add(df.loc[i,worddic[nowcommend]] +care)
                            #--------------------------------------------------------------------------------
                            elif nowcommend == '聯絡電話':
                                if cell.text not in testset:
                                    homnnum,clphone = cell.text.split('\n')
                                    homnnum = homnnum.replace(' ','')
                                    clphone = clphone.replace(' ','')
                                    homnnum = homnnum + str(df.loc[i,'電話(公)'])
                                    clphone = clphone + '0' +str(df.loc[i,'電話(行動)'])
                                    cell.text = homnnum + '\n' + clphone
                                    testset.add(cell.text)
                            elif nowcommend =='上課別':
                                index = df.loc[i,str(worddic[nowcommend])]
                                reCell  = table.cell(idxr+1 , idxc-2) 
                                reCell.text  = stu[int(index)]
                            elif nowcommend =='年級':
                                reCell  = table.cell(idxr+1 , idxc-1)
                                reCell.text = str(df.loc[i,'年級'])
                            elif nowcommend =='班別': 
                                reCell  = table.cell(idxr+1 , idxc-1)
                                reCell.text = str(df.loc[i, '班別'])
                            elif nowcommend =='座號':
                                reCell  = table.cell(idxr ,idxc+6)
                                number = str(df_print.loc[i,'座號' ])
                                reCell.text = number
                            elif nowcommend =='學制':
                                school_type_id = df_print.loc[i,'學制']
                                school_type = study_type_list[school_type_id]
                                cell.text = cell.text.replace('\t','')
                                checkboxlst = list(map(str,cell.text.split('\n')))
                                option_lst = []
                                cell.text = ''
                                paragraph = cell.paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                option_len = 7
                                # -----------------------------------------------------------------------------
                                for idx, j in enumerate(checkboxlst):
                                    times = j.count('□')
                                    temlst = []
                                    if times == 0 :
                                        underline_option +=j +'\n'
                                        option_lst.append([j,'\n'])
                                    else:
                                        for k in range(times):
                                            option_not_edit = j[len(j)-j[::-1].index('□')-1:]
                                            squard , option = option_not_edit[:1],option_not_edit[1:]
                                            if k == 0 and idx != len(checkboxlst)-1:
                                                option +='\n'
                                            else:
                                                add_space = option_len - len(option)
                                                if add_space <0:
                                                    add_space = 0
                                                option += '　'*(add_space)
                                            temlst.append([squard,option])
                                            j = j.replace(option_not_edit,'')
                                    temlst = temlst[::-1]
                                    for x in temlst :
                                        option_lst.append(x)
                                for option in option_lst:
                                    if school_type in option[1]:
                                        option[0] = '■'
                                        break
                                #---------------------------------------------------------------------------------
                                response = ''
                                for res in option_lst:
                                    run = paragraph.add_run(res[0])
                                    run.font.size = Pt(12)
                                    run.font.name = '標楷體'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                    run2 = paragraph.add_run(res[1])
                                    run2.font.size = Pt(12)
                                    run2.font.name ='Times New Roman'  
                                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')     
                                    rPr = run2._element.get_or_add_rPr()
                                    spacing = OxmlElement('w:spacing')
                                    spacing.set(qn('w:val'), '0.5')  
                                    rPr.append(spacing)
                            elif nowcommend == '身分別(一般報檢人免填)':
                                idCode =str(df.loc[i,'特定對象身份別'])
                                if idCode != 'nan':
                                    id = dfsg.loc[float(idCode), 'Unnamed: 21']
                                    checkboxlst = list(map(str,cell.text.split('\n')))
                                    cell.text = ''
                                    option_lst = [] 
                                    paragraph = cell.paragraphs[0]
                                    paragraph_format = paragraph.paragraph_format
                                    paragraph_format.line_spacing = Pt(9)
                                    underline_option = ''
                                    for idx, j in enumerate(checkboxlst):
                                        times = j.count('□')
                                        temlst = []
                                        if times == 0 :
                                            underline_option +=j +'\n'
                                            option_lst.append([j,'\n'])
                                        else:
                                            for k in range(times):
                                                option_not_edit = j[len(j)-j[::-1].index('□')-1:]
                                                squard , option = option_not_edit[:1],option_not_edit[1:]
                                                if k == 0 and idx != len(checkboxlst)-1:
                                                    option +='\n'
                                                temlst.append([squard,option])
                                                j = j.replace(option_not_edit,'')
                                        temlst = temlst[::-1]
                                        for x in temlst :
                                            option_lst.append(x)
                                    for option in option_lst:
                                        if id in option[1]:
                                            option[0] = '■'
                                            break
                                    response = ''
                                    for res in option_lst:
                                        run = paragraph.add_run(res[0])
                                        run.font.size = Pt(7.8)
                                        run.font.name = '標楷體'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                        run2 = paragraph.add_run(res[1])
                                        run2.font.size = Pt(7.8)
                                        run2.font.name ='Times New Roman'  
                                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')     
                                        if '限以上身分，需另填寫補助申請書，不申請補助者免附' in res[0]:
                                            run3 = paragraph.add_run('————————————————————\n') 
                                            run3.font.name ='Times New Roman'  
                                            run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')             
                                            rPr2 = run3._element.get_or_add_rPr()
                                            spacing = OxmlElement('w:spacing')
                                            spacing.set(qn('w:val'), '0.5')  
                                            rPr2.append(spacing)
                                        rPr = run2._element.get_or_add_rPr()
                                        spacing = OxmlElement('w:spacing')
                                        spacing.set(qn('w:val'), '0.5')  
                                        rPr.append(spacing)
                            elif nowcommend =='報檢職類':
                                type_dict = {
                                    '視覺':'視覺傳達設計',
                                    '會計人工':'會計事務 -人工記帳',
                                    '會資':'會計事務 -資訊',
                                    '門市':'門市服務'
                                }
                                id_type = int(df_print.loc[i,'測驗類別'])
                                test_type = test_type_lst[id_type]
                                temstr = test_type[::-1]
                                temstr = temstr[:2]
                                test_type = test_type.replace(temstr[::-1],'')
                                test_type = type_dict[test_type]
                                checkboxlst = list(map(str,cell.text.split('\n')))
                                option_lst = []
                                cell.text = ''
                                paragraph = cell.paragraphs[0]
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.line_spacing = Pt(12)  
                            # ---------------------------------------------------------------------------------
                            # 分割checkbox 與選項內容
                                for idx,j in enumerate(checkboxlst):
                                        times = j.count('□')
                                        temlst = []
                                        if times == 0 :
                                            underline_option +=j +'\n' 
                                            option_lst.append([j,'\n'])
                                        else:
                                            for k in range(times):
                                                option_not_edit = j[len(j)-j[::-1].index('□')-1:]
                                                squard , option = option_not_edit[:1],option_not_edit[1:]
                                                if k == 0 and idx != (len(checkboxlst)-1):
                                                    option +='\n'
                                                temlst.append([squard,option])
                                                j = j.replace(option_not_edit,'')
                                        temlst = temlst[::-1]
                                        for x in temlst :
                                            option_lst.append(x)
                                for res in option_lst:
                                    if test_type in res[1]:
                                        res[0] = '■'
                                        break
                                for res in option_lst:
                                        if '會計事務 -資訊' in res[1]:
                                            res[1] = res[1][:len(res[1])-1]
                                        run = paragraph.add_run(res[0])
                                        run.font.size = Pt(10.5)
                                        run.font.name = '標楷體'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                        run2 = paragraph.add_run(res[1])
                                        run2.font.size = Pt(10.5)
                                        run2.font.name ='Times New Roman'  
                                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')  
                                        rPr = run._element.get_or_add_rPr()
                                        spacing = OxmlElement('w:spacing')
                                        spacing.set(qn('w:val'), '0')  
                                        rPr.append(spacing)
                                        rPr = run2._element.get_or_add_rPr()
                                        spacing = OxmlElement('w:spacing')
                                        spacing.set(qn('w:val'), '0')  
                                        rPr.append(spacing) 
                            # ---------------------------------------------------------------------------------
                            elif '實貼身分證【正面】' in nowcommend:
                                cell.text = ''
                                class_id = '0'
                                class_id += str(df_print.loc[i,'學號'])
                                paragraph = cell.paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run()
                                path =f"./excel_to_word/processingData/{class_id}.jpg" 
                                run.add_picture(path)
                            elif   '檢定區別' in nowcommend:
                                test_type = test_type_lst[int(df_print.loc[i,'測驗類別'])]
                                test_type_type  = test_type[::-1][:2][::-1]
                                option_lst=[]
                                checkboxlst = list(map(str,cell.text.split('\n')))
                                cell.text = ''
                                paragraph = cell.paragraphs[0]
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.line_spacing = Pt(12)
                                test_type_dic = {
                                    '全測':'學術科全測',
                                    '免術':'免試術科',
                                    '免學':'免試學科'
                                }
                                for idx,j in enumerate(checkboxlst):
                                        times = j.count('□')
                                        temlst = []
                                        if times == 0 :
                                            underline_option +=j +'\n' 
                                            option_lst.append([j,'\n'])
                                        else:
                                            for k in range(times):
                                                option_not_edit = j[len(j)-j[::-1].index('□')-1:]
                                                squard , option = option_not_edit[:1],option_not_edit[1:]
                                                other = option[5:]
                                                if other == ' ':
                                                    if k == 0 and idx != (len(checkboxlst)-1):
                                                        option +='\n'
                                                    temlst.append([squard,option])
                                                    j = j.replace(option_not_edit,'')
                                                else:
                                                    option = option.replace(other,'')
                                                    if k == 0 and idx != (len(checkboxlst)-1):
                                                        other +='\n'
                                                    temlst.append([squard,option,other])
                                                    j = j.replace(option_not_edit,'')
                                        temlst = temlst[::-1]
                                        for x in temlst :
                                            option_lst.append(x)
                                for res in option_lst:
                                    if test_type_dic[test_type_type] in res[1]:
                                        res[0] = '■'
                                        break
                                #---------------------------------------------------------------------------------------------------
                                for res in option_lst:
                                    run = paragraph.add_run(res[0])
                                    run.font.size = Pt(12)
                                    run.font.name = '標楷體'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                    run2 = paragraph.add_run(res[1])
                                    run2.font.size = Pt(12)
                                    run2.font.name = 'Times New Roman'
                                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                    if len(res)>2:
                                        run3 = paragraph.add_run(res[2])
                                        run3.font.size = Pt(8)
                                        run3.font.name = 'Times New Roman'
                                        run3._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                    rPr = run2._element.get_or_add_rPr()
                                    spacing = OxmlElement('w:spacing')
                                    spacing.set(qn('w:val'), '0.5')  
                                    rPr.append(spacing)
                            testset.add(cell.text)  
                            nowcommend = cell.text
                            nowcommend = nowcommend.replace('\n','')
                    else:
                            try:
                                if nowcommend =='就讀學校':
                                    school  = df.loc[i,worddic[nowcommend]]
                                    cell.text = dfsc.loc[str(school),'Unnamed: 9']
                                elif nowcommend =='就讀科系':
                                    sub = df.loc[i,worddic[nowcommend]]
                                    cell.text=subject[int(sub)]
                                else:
                                    report = str(df.loc[i,worddic[nowcommend]])
                                    if report =='nan':
                                        cell.text = ''
                                    else:
                                        cell.text = report 
                            except:
                                break
            new_file_path = school_id+'.docx'
            doc.save('./excel_to_word/alreadyPDF/'+new_file_path) 
            convert('./excel_to_word/alreadyPDF/'+new_file_path)
            merger.append('./excel_to_word/./alreadyPDF/'+school_id+'.pdf')
        merger.write("./excel_to_word/alreadyPDF/result.pdf")
        merger.close()
        files = glob.glob('./excel_to_word/processingData/*.*')
        for file in files:
            os.remove(file)
        files = glob.glob('./excel_to_word/alreadyPDF/*.*')
        for file in files:
            if  'result.pdf' in file:
                continue
            os.remove(file)
        return True
    except:
        return False
    